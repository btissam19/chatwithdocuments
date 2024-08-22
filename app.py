import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from docx import Document
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=api_key)

# Function to extract text from PDFs
def get_pdf_text(pdf_docs):
    text = ""
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

# Function to get text chunks
def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = text_splitter.split_text(text)
    return chunks

# Function to create and save the vector store
def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")

# Function to load the conversational chain
def get_conversational_chain():
    prompt_template = """
    Answer the question as detailed as possible from the provided context, make sure to provide all the details, if the answer is not in
    provided context just say, "answer is not available in the documents you upload", don't provide the wrong answer\n\n
    Context:\n {context}?\n
    Question: \n{question}\n

    Answer:
    """

    model = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

    return chain

# Function to handle user input and get the response
def user_input(user_question):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    
    docs = new_db.similarity_search(user_question)
    print(f"Docs retrieved: {docs}")  # Debugging statement
    
    chain = get_conversational_chain()
    response = chain(
        {"input_documents": docs, "question": user_question}, 
        return_only_outputs=True
    )
    
    print(f"Response: {response}")  # Debugging statement
    return response["output_text"]

# Function to create and save a DOCX report
def create_docx_report(chat_history, file_name="report.docx"):
    doc = Document()
    doc.add_heading('Chat History Report', 0)

    for message in chat_history:
        if message['role'] == 'user':
            doc.add_heading('You:', level=1)
            doc.add_paragraph(message['text'])
        else:
            doc.add_heading('Assistant:', level=1)
            doc.add_paragraph(message['text'])

    doc.save(file_name)
    return file_name

def main():
    st.set_page_config(page_title="Chat with Documents")
    st.header("Chat with Various Pdf Documents and generate your custum  report ")
    # Initialize chat history in session state
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []

    # Display chat history
    for message in st.session_state.chat_history:
        if message['role'] == 'user':
            st.write(f"*You:* {message['text']}")
        else:
            st.write(f"*Assistant:* {message['text']}")

    # Input for new questions
    user_question = st.text_input("Ask a Question from the Files")

    if user_question:
        response = user_input(user_question)
        st.session_state.chat_history.append({"role": "user", "text": user_question})
        st.session_state.chat_history.append({"role": "assistant", "text": response})
        st.write("*Assistant:*", response)  # Ensure the response is shown

    with st.sidebar:
        st.title("Menu:")
        uploaded_files = st.file_uploader("Upload your Files and Click on the Submit & Process Button", accept_multiple_files=True, type=['pdf'])
        if st.button("Submit & Process"):
            with st.spinner("Processing..."):
                raw_text = ""
                pdf_docs = [file for file in uploaded_files if file.name.endswith('.pdf')]
                
                if pdf_docs:
                    raw_text += get_pdf_text(pdf_docs)
                
                text_chunks = get_text_chunks(raw_text)
                get_vector_store(text_chunks)
                st.success("Documents processed and indexed.")

        if st.button("Generate DOCX Report"):
            if st.session_state.chat_history:
                file_name = create_docx_report(st.session_state.chat_history)
                with open(file_name, "rb") as file:
                    st.download_button(
                        label="Download DOCX Report",
                        data=file,
                        file_name=file_name
                    )

if __name__ == "__main__":
    main()